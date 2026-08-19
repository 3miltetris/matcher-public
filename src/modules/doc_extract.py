"""
doc_extract.py — streamlit-free document text extraction.

Shared by the Drive Sync view and jobs/drive_sync_job.py. Extraction
cascade lifted from views/resume_importer.py (the views stay top-level
Streamlit scripts and cannot be imported).

Supported: PDF (fitz), DOCX (ZIP XML → python-docx → fitz cascade),
XLSX (per-sheet CSV dump), TXT/CSV/MD (plain decode).
"""

import io
import zipfile
from xml.etree import ElementTree as ET

import fitz  # PyMuPDF
import pandas as pd

TEXT_LIMIT   = 30_000   # chars per document
MIN_TEXT_LEN = 100      # below this, extraction is treated as failed


def detect_file_type(content: bytes, name: str, mime: str = '') -> str:
    """Return 'pdf' | 'docx' | 'xlsx' | 'txt' | 'csv' | 'unknown'."""
    lname = (name or '').lower()
    lmime = (mime or '').lower()
    if 'pdf' in lmime or lname.endswith('.pdf') or (len(content) >= 4 and content[:4] == b'%PDF'):
        return 'pdf'
    if 'spreadsheetml' in lmime or lname.endswith('.xlsx'):
        return 'xlsx'
    if ('wordprocessingml' in lmime or lname.endswith(('.docx', '.doc'))
            or (len(content) >= 2 and content[:2] == b'PK')):
        return 'docx'
    if lname.endswith('.csv') or 'text/csv' in lmime:
        return 'csv'
    if lname.endswith(('.txt', '.md')) or lmime.startswith('text/'):
        return 'txt'
    return 'unknown'


def extract_pdf(content: bytes, max_pages: int = 100) -> str:
    try:
        doc   = fitz.open(stream=content, filetype='pdf')
        parts = []
        for i, page in enumerate(doc):
            if i >= max_pages:
                break
            text = page.get_text('text').strip()
            if text:
                parts.append(text)
        doc.close()
        return ' '.join(parts)[:TEXT_LIMIT]
    except Exception:
        return ''


def extract_docx_xml(content: bytes) -> str:
    """Parse word/document.xml directly to pull every <w:t> text run.
    Captures paragraphs, table cells, AND text boxes — the python-docx API
    misses text boxes."""
    try:
        with zipfile.ZipFile(io.BytesIO(content)) as z:
            with z.open('word/document.xml') as f:
                xml_bytes = f.read()
        ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        root  = ET.fromstring(xml_bytes)
        texts = [el.text for el in root.findall('.//w:t', ns) if el.text]
        return ' '.join(texts)[:TEXT_LIMIT]
    except Exception:
        return ''


def extract_docx(content: bytes) -> str:
    # Primary: XML parse — the only method that sees text boxes
    text = extract_docx_xml(content)
    if len(text.strip()) >= MIN_TEXT_LEN:
        return text
    # Fallback: python-docx paragraphs + table cells
    try:
        import docx as docx_lib
        doc   = docx_lib.Document(io.BytesIO(content))
        parts = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    parts.append(cell.text)
        text = ' '.join(p for p in parts if p.strip())[:TEXT_LIMIT]
        if len(text.strip()) >= MIN_TEXT_LEN:
            return text
    except Exception:
        pass
    # Last resort: PyMuPDF
    try:
        doc  = fitz.open(stream=content, filetype='docx')
        text = ' '.join(page.get_text() for page in doc)[:TEXT_LIMIT]
        if len(text.strip()) >= MIN_TEXT_LEN:
            return text
    except Exception:
        pass
    return ''


def extract_xlsx(content: bytes) -> str:
    try:
        sheets = pd.read_excel(io.BytesIO(content), sheet_name=None,
                               dtype=str, engine='openpyxl')
        parts = []
        for sheet_name, df in sheets.items():
            df = df.dropna(how='all')
            if df.empty:
                continue
            parts.append(f'## Sheet: {sheet_name}\n{df.to_csv(index=False)}')
            if sum(len(p) for p in parts) > TEXT_LIMIT:
                break
        return '\n\n'.join(parts)[:TEXT_LIMIT]
    except Exception:
        return ''


def _decode_text(content: bytes) -> str:
    try:
        return content.decode('utf-8')[:TEXT_LIMIT]
    except UnicodeDecodeError:
        try:
            return content.decode('latin-1')[:TEXT_LIMIT]
        except Exception:
            return ''


def extract_text(content: bytes, name: str, mime: str = '') -> tuple[str, str]:
    """Dispatch on detected type. Returns (text, file_type); text is '' on failure."""
    ftype = detect_file_type(content, name, mime)
    if ftype == 'pdf':
        return extract_pdf(content), 'pdf'
    if ftype == 'docx':
        return extract_docx(content), 'docx'
    if ftype == 'xlsx':
        return extract_xlsx(content), 'xlsx'
    if ftype in ('txt', 'csv'):
        return _decode_text(content), ftype
    return '', 'unknown'
