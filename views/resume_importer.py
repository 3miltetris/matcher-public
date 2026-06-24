"""
Resume Importer
---------------
Upload a HubSpot contacts CSV that includes a resume URL column.
Fetches each resume (PDF or DOCX) from HubSpot, extracts text, generates a
GPT expertise summary, embeds with text-embedding-ada-002, and saves to GCS
under data/resumes/{YYYY-MM-DD}_{hex6}.parquet.
"""

import io
import secrets
import uuid
from concurrent.futures import ThreadPoolExecutor
from concurrent.futures import as_completed as futures_completed
from datetime import datetime

import fitz  # pymupdf
import pandas as pd
import requests as req_lib
import streamlit as st
from google.cloud import storage
from google.oauth2 import service_account
from openai import OpenAI

from src.modules.Embedding.text_embedder import TextProcessor

# ── Constants ──────────────────────────────────────────────────────────────

_BUCKET         = 'cc-matcher-bucket-jeg-v1'
_RESUMES_PREFIX = 'data/resumes/'
_TEXT_LIMIT     = 12_000  # chars of resume text fed to GPT
_MIN_TEXT_LEN   = 400     # minimum chars to consider extraction successful

_SUMMARISE_SYSTEM = (
    'You are analyzing resume text. Write a 3-5 sentence professional summary covering: '
    '(1) primary technical skills and tools, '
    '(2) domain and industry expertise, '
    '(3) years of relevant experience, and '
    '(4) notable project types or accomplishments. '
    'CRITICAL: only use information explicitly stated in the text. '
    'Do NOT infer, assume, or invent any detail not present. '
    'If the text does not contain enough information to write a grounded summary, '
    'respond with exactly: INSUFFICIENT_TEXT'
)

_FIELD_CANDIDATES: dict[str, list[str]] = {
    'email':      ['email', 'email address', 'work email', 'e-mail'],
    'resume_url': ['resume', 'resume url', 'resume_url', 'cv', 'cv url', 'file', 'file url',
                   'attachment', 'resume link', 'document'],
    'firstName':  ['first name', 'firstname', 'first_name', 'given name'],
    'lastName':   ['last name', 'lastname', 'last_name', 'surname', 'family name'],
    'phone':      ['phone', 'phone number', 'mobile', 'telephone', 'cell'],
    'company':    ['company', 'company name', 'companyname', 'organization', 'employer', 'account name'],
}


def _detect_col(columns: list[str], field: str) -> str | None:
    lower = {c.lower().strip(): c for c in columns}
    for candidate in _FIELD_CANDIDATES[field]:
        if candidate in lower:
            return lower[candidate]
    return None


# ── GCS ────────────────────────────────────────────────────────────────────

def _get_storage_client() -> storage.Client:
    creds = service_account.Credentials.from_service_account_info(
        st.secrets['gcp_service_account']
    )
    return storage.Client(credentials=creds)


def _load_existing_emails(client: storage.Client) -> set[str]:
    blobs = client.list_blobs(_BUCKET, prefix=_RESUMES_PREFIX)
    emails: set[str] = set()
    for blob in blobs:
        if not blob.name.endswith('.parquet'):
            continue
        try:
            df = pd.read_parquet(io.BytesIO(blob.download_as_bytes()), columns=['email'])
            emails.update(df['email'].dropna().str.strip().str.lower().tolist())
        except Exception:
            pass
    return emails


# ── Fetch helpers ──────────────────────────────────────────────────────────

def _fetch_resume(url: str, hs_token: str | None) -> tuple[bytes | None, str, str]:
    """Return (content_bytes, content_type, error_detail).
    error_detail is non-empty only on failure and describes why."""
    headers = {'User-Agent': 'MatcherBot/1.0'}
    if hs_token:
        headers['Authorization'] = f'Bearer {hs_token}'
    try:
        resp = req_lib.get(url, headers=headers, timeout=25, allow_redirects=True)
        if resp.status_code == 200:
            ct = resp.headers.get('Content-Type', '')
            if 'text/html' in ct:
                return None, '', f'200 OK but Content-Type is text/html (auth redirect?)'
            return resp.content, ct, ''
        return None, '', f'HTTP {resp.status_code}'
    except Exception as exc:
        return None, '', f'exception: {exc}'


def _detect_file_type(content: bytes, url: str, content_type: str) -> str:
    ct = content_type.lower()
    if 'pdf' in ct or url.lower().endswith('.pdf') or (len(content) >= 4 and content[:4] == b'%PDF'):
        return 'pdf'
    if ('wordprocessingml' in ct or url.lower().endswith('.docx')
            or (len(content) >= 2 and content[:2] == b'PK')):
        return 'docx'
    return 'unknown'


def _extract_docx_xml(content: bytes) -> str:
    """Parse word/document.xml directly to pull every <w:t> text run.
    Captures paragraphs, table cells, AND text boxes — the python-docx API misses text boxes,
    which most resume templates use for their column/sidebar layout."""
    import zipfile
    from xml.etree import ElementTree as ET
    try:
        with zipfile.ZipFile(io.BytesIO(content)) as z:
            with z.open('word/document.xml') as f:
                xml_bytes = f.read()
        ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        root = ET.fromstring(xml_bytes)
        texts = [el.text for el in root.findall('.//w:t', ns) if el.text]
        return ' '.join(texts)[:_TEXT_LIMIT]
    except Exception:
        return ''


def _extract_text(content: bytes, url: str, content_type: str) -> tuple[str, str]:
    ftype = _detect_file_type(content, url, content_type)
    if ftype == 'pdf':
        try:
            doc   = fitz.open(stream=content, filetype='pdf')
            parts = []
            for page in doc:
                # get_text("text") is reliable for native PDFs;
                # iterating blocks catches some layouts that the flat join misses
                text = page.get_text("text").strip()
                if text:
                    parts.append(text)
            return ' '.join(parts)[:_TEXT_LIMIT], 'pdf'
        except Exception:
            return '', 'pdf'
    if ftype == 'docx':
        # Primary: XML parse — the only method that sees text boxes
        text = _extract_docx_xml(content)
        if len(text.strip()) >= _MIN_TEXT_LEN:
            return text, 'docx'
        # Fallback: python-docx paragraphs + table cells
        try:
            import docx as docx_lib
            doc   = docx_lib.Document(io.BytesIO(content))
            parts = [p.text for p in doc.paragraphs]
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        parts.append(cell.text)
            text = ' '.join(p for p in parts if p.strip())[:_TEXT_LIMIT]
            if len(text.strip()) >= _MIN_TEXT_LEN:
                return text, 'docx'
        except Exception:
            pass
        # Last resort: PyMuPDF
        try:
            doc = fitz.open(stream=content, filetype='docx')
            text = ' '.join(page.get_text() for page in doc)[:_TEXT_LIMIT]
            if len(text.strip()) >= _MIN_TEXT_LEN:
                return text, 'docx'
        except Exception:
            pass
        return '', 'docx'
    return '', 'unknown'


def _run_fetch_and_extract(rows: list[dict], hs_token: str | None, progress) -> list[dict]:
    results: list[dict] = [{}] * len(rows)
    total = len(rows)

    def _one(idx: int, row: dict) -> tuple[int, dict]:
        url = str(row.get('resume_url') or '').strip()
        if not url:
            return idx, {**row, 'resume_text': '', 'file_type': 'missing', 'fetch_error': 'no url'}
        content, ct, err = _fetch_resume(url, hs_token)
        if content is None:
            return idx, {**row, 'resume_text': '', 'file_type': 'fetch_failed', 'fetch_error': err}
        text, ftype = _extract_text(content, url, ct)
        return idx, {**row, 'resume_text': text, 'file_type': ftype, 'fetch_error': ''}

    with ThreadPoolExecutor(max_workers=8) as pool:
        futures = {pool.submit(_one, i, r): i for i, r in enumerate(rows)}
        done = 0
        for future in futures_completed(futures):
            idx, result = future.result()
            results[idx] = result
            done += 1
            progress.progress(done / total, text=f'Fetching resumes… {done}/{total}')

    return results


# ── Summarisation ──────────────────────────────────────────────────────────

def _run_summarization(rows: list[dict], openai_key: str, progress) -> list[dict]:
    results: list[dict] = [{}] * len(rows)
    total = len(rows)
    client = OpenAI(api_key=openai_key)

    def _one(idx: int, row: dict) -> tuple[int, dict]:
        text = str(row.get('resume_text', '')).strip()
        if len(text) < _MIN_TEXT_LEN:
            return idx, {**row, 'expertise_summary': ''}
        try:
            resp = client.chat.completions.create(
                model='gpt-3.5-turbo',
                max_tokens=400,
                messages=[
                    {'role': 'system', 'content': _SUMMARISE_SYSTEM},
                    {'role': 'user',   'content': text},
                ],
            )
            summary = resp.choices[0].message.content.strip()
            if summary == 'INSUFFICIENT_TEXT':
                summary = ''
            return idx, {**row, 'expertise_summary': summary}
        except Exception as e:
            return idx, {**row, 'expertise_summary': f'SUMMARY_ERROR: {e}'}

    with ThreadPoolExecutor(max_workers=10) as pool:
        futures = {pool.submit(_one, i, r): i for i, r in enumerate(rows)}
        done = 0
        for future in futures_completed(futures):
            idx, result = future.result()
            results[idx] = result
            done += 1
            progress.progress(done / total, text=f'Summarising… {done}/{total}')

    return results


# ── Session state ──────────────────────────────────────────────────────────

for _k in ('ri_raw_df', 'ri_mapped_rows', 'ri_deduped_rows', 'ri_processed_rows',
           'ri_email_col', 'ri_url_col'):
    if _k not in st.session_state:
        st.session_state[_k] = None


# ── Page ───────────────────────────────────────────────────────────────────

st.title('📄 Resume Importer')
st.caption(
    'Upload a HubSpot contacts export with a resume URL column → fetch + extract text → '
    'GPT expertise summary → embed → save to GCS.'
)

# ── 1 · Upload ─────────────────────────────────────────────────────────────

st.subheader('1 · Upload HubSpot CSV')
uploaded = st.file_uploader(
    'Upload CSV or Excel',
    type=['csv', 'xlsx', 'xls'],
    label_visibility='collapsed',
)

if uploaded:
    try:
        if uploaded.name.endswith(('.xlsx', '.xls')):
            raw = pd.read_excel(uploaded, dtype=str)
        else:
            try:
                raw = pd.read_csv(uploaded, dtype=str, encoding='utf-8')
            except UnicodeDecodeError:
                uploaded.seek(0)
                raw = pd.read_csv(uploaded, dtype=str, encoding='latin-1')
        raw = raw.dropna(how='all')

        if (st.session_state.ri_raw_df is None
                or len(raw) != len(st.session_state.ri_raw_df)):
            st.session_state.ri_raw_df        = raw
            st.session_state.ri_mapped_rows   = None
            st.session_state.ri_deduped_rows  = None
            st.session_state.ri_processed_rows = None
    except Exception as e:
        st.error(f'Could not read file: {e}')

if st.session_state.ri_raw_df is None:
    st.stop()

df_raw = st.session_state.ri_raw_df
st.caption(f'**{len(df_raw):,}** rows loaded.')
st.dataframe(df_raw.head(5), hide_index=True, use_container_width=True)

# ── 2 · Column mapping ─────────────────────────────────────────────────────

st.divider()
st.subheader('2 · Map columns')

cols     = df_raw.columns.tolist()
none_opt = '— none —'
col_opts = [none_opt] + cols


def _sel(field: str, label: str, required: bool = False) -> str | None:
    detected = _detect_col(cols, field)
    idx      = col_opts.index(detected) if detected in col_opts else 0
    val      = st.selectbox(
        label + (' *' if required else ''),
        col_opts,
        index=idx,
        key=f'ri_map_{field}',
    )
    return val if val != none_opt else None


mc1, mc2, mc3 = st.columns(3)
with mc1:
    m_email = _sel('email',      'Email',      required=True)
    m_url   = _sel('resume_url', 'Resume URL', required=True)
with mc2:
    m_first = _sel('firstName', 'First name')
    m_last  = _sel('lastName',  'Last name')
with mc3:
    m_phone = _sel('phone',   'Phone')
    m_comp  = _sel('company', 'Company')

if not m_email:
    st.warning('Email column is required.')
    st.stop()
if not m_url:
    st.warning('Resume URL column is required.')
    st.stop()

# Invalidate downstream state when key columns change
if (st.session_state.ri_email_col is not None and (
        st.session_state.ri_email_col != m_email
        or st.session_state.ri_url_col != m_url)):
    st.session_state.ri_mapped_rows   = None
    st.session_state.ri_deduped_rows  = None
    st.session_state.ri_processed_rows = None


def _build_mapped_rows() -> list[dict]:
    def _pick(col: str | None, row: pd.Series) -> str:
        return str(row[col]).strip() if col and col in row.index else ''

    rows = []
    for _, row in df_raw.iterrows():
        email = _pick(m_email, row).lower()
        url   = _pick(m_url,   row)
        if not email or not url:
            continue
        rows.append({
            'email':      email,
            'resume_url': url,
            'firstName':  _pick(m_first, row),
            'lastName':   _pick(m_last,  row),
            'phone':      _pick(m_phone, row),
            'company':    _pick(m_comp,  row),
        })
    return rows


# ── 3 · Deduplicate ────────────────────────────────────────────────────────

st.divider()
st.subheader('3 · Deduplicate')

if st.button('🔍 Check for duplicates', key='ri_dedup_btn'):
    with st.spinner('Checking GCS for existing emails…'):
        try:
            mapped = _build_mapped_rows()
            client  = _get_storage_client()
            existing = _load_existing_emails(client)
            new_rows = [r for r in mapped if r['email'] not in existing]
            st.session_state.ri_mapped_rows   = mapped
            st.session_state.ri_deduped_rows  = new_rows
            st.session_state.ri_email_col     = m_email
            st.session_state.ri_url_col       = m_url
            st.session_state.ri_processed_rows = None
            st.rerun()
        except Exception as e:
            st.error(f'Dedup check failed: {e}')

if st.session_state.ri_deduped_rows is None:
    mapped_preview = _build_mapped_rows()
    st.caption(f'**{len(mapped_preview):,}** rows with email + URL. Click to check for duplicates.')
    st.stop()

deduped = st.session_state.ri_deduped_rows
n_all   = len(st.session_state.ri_mapped_rows or [])
n_exist = n_all - len(deduped)

dm1, dm2, dm3 = st.columns(3)
dm1.metric('Valid rows',     f'{n_all:,}')
dm2.metric('Already stored', f'{n_exist:,}')
dm3.metric('New to import',  f'{len(deduped):,}')

if not deduped:
    st.success('All contacts already stored — nothing to import.')
    st.stop()

# ── 4 · Fetch + Summarise ──────────────────────────────────────────────────

st.divider()
st.subheader('4 · Fetch resumes & generate expertise summaries')

hs_token = st.secrets.get('hubspot_api_key', None)

if st.session_state.ri_processed_rows is None:
    if st.button('▶ Fetch & summarise', type='primary'):
        oai_key = st.secrets['openai_api_key']

        prog_fetch = st.progress(0, text='Fetching resumes…')
        fetched    = _run_fetch_and_extract(deduped, hs_token, prog_fetch)
        prog_fetch.empty()

        statuses: dict[str, int] = {}
        for r in fetched:
            ft = r.get('file_type', 'unknown')
            statuses[ft] = statuses.get(ft, 0) + 1
        st.write('**Fetch results:**', statuses)

        # Show a sample of failure reasons to diagnose auth/URL issues
        failed_rows = [r for r in fetched if r.get('file_type') == 'fetch_failed']
        if failed_rows:
            with st.expander(f'Failure diagnostics (sample of {min(5, len(failed_rows))})'):
                for r in failed_rows[:5]:
                    st.markdown(f'**URL:** `{r.get("resume_url", "")[:120]}`')
                    st.markdown(f'**Reason:** `{r.get("fetch_error", "unknown")}`')
                    st.divider()

        ok_rows = [r for r in fetched
                   if len(str(r.get('resume_text', '')).strip()) >= _MIN_TEXT_LEN]
        failed  = len(fetched) - len(ok_rows)
        if failed:
            st.warning(f'{failed} resume(s) could not be fetched or had no extractable text.')
        if not ok_rows:
            st.error('No resume text extracted — check URLs and try again.')
            st.stop()

        with st.expander(f'Preview extracted text ({len(ok_rows)} resumes) — verify before summarising'):
            for r in ok_rows[:5]:
                name = f"{r.get('firstName', '')} {r.get('lastName', '')}".strip() or r.get('email', '?')
                st.markdown(f'**{name}** (`{r.get("file_type")}`)')
                st.text(str(r.get('resume_text', ''))[:600])
                st.divider()
            if len(ok_rows) > 5:
                st.caption(f'Showing 5 of {len(ok_rows)}.')

        prog_sum   = st.progress(0, text='Summarising…')
        summarised = _run_summarization(ok_rows, oai_key, prog_sum)
        prog_sum.empty()

        st.session_state.ri_processed_rows = summarised
        st.rerun()
    st.stop()  # wait for the button click before rendering steps 5+
else:
    summarised = st.session_state.ri_processed_rows

proc_df = pd.DataFrame(summarised)
preview_cols = [c for c in ['email', 'firstName', 'lastName', 'file_type', 'expertise_summary']
                if c in proc_df.columns]
st.dataframe(proc_df[preview_cols].head(20), use_container_width=True, hide_index=True)
if len(proc_df) > 20:
    st.caption(f'Showing 20 of {len(proc_df)} rows.')

# ── 5 · Embed & save ───────────────────────────────────────────────────────

st.divider()
st.subheader('5 · Embed & save to GCS')

if st.button('▶ Embed & save', type='primary'):
    oai_key  = st.secrets['openai_api_key']
    tp       = TextProcessor(api_key=oai_key)
    embed_df = proc_df.copy()

    prog = st.progress(0, text='Embedding…')
    embeddings = []
    total = len(embed_df)
    for i, (_, row) in enumerate(embed_df.iterrows()):
        text = str(row.get('expertise_summary') or '').strip()
        embeddings.append(tp.get_embedding(text) if text else [])
        prog.progress((i + 1) / total, text=f'Embedding… {i + 1}/{total}')
    prog.empty()

    embed_df['embeddings']   = embeddings
    embed_df['processed_at'] = datetime.utcnow().date().isoformat()
    embed_df['uuid']         = [str(uuid.uuid4()) for _ in range(len(embed_df))]
    embed_df = embed_df.drop(columns=['resume_text'], errors='ignore')

    ordered_cols = [c for c in [
        'uuid', 'email', 'firstName', 'lastName', 'phone', 'company',
        'resume_url', 'file_type', 'expertise_summary', 'embeddings', 'processed_at',
    ] if c in embed_df.columns]
    out_df = embed_df[ordered_cols]

    date_str  = datetime.utcnow().strftime('%Y-%m-%d')
    hex_str   = secrets.token_hex(3)
    blob_path = f'{_RESUMES_PREFIX}resumes_{date_str}_{hex_str}.parquet'

    try:
        client = _get_storage_client()
        buf    = io.BytesIO()
        out_df.to_parquet(buf, index=False)
        buf.seek(0)
        blob = client.bucket(_BUCKET).blob(blob_path)
        blob.upload_from_file(buf, content_type='application/octet-stream')
        st.success(f'Saved {len(out_df):,} resumes → `{blob_path}`')
        for _k in ('ri_raw_df', 'ri_mapped_rows', 'ri_deduped_rows', 'ri_processed_rows',
                   'ri_email_col', 'ri_url_col'):
            st.session_state[_k] = None
    except Exception as e:
        st.error(f'GCS upload failed: {e}')
